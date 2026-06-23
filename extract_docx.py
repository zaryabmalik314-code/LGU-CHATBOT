"""
Finds .docx URLs inside lgu_urls_clean.txt (they got mixed in there),
downloads and extracts text from them.
Saves results as JSONL in lgu_docx_content.jsonl
"""

import urllib.request
import json
import time
import io
from datetime import datetime
import docx  # from python-docx package

INPUT_FILE = "lgu_urls_clean.txt"
OUTPUT_FILE = "lgu_docx_content.jsonl"
FAILED_FILE = "lgu_docx_failed.txt"

HEADERS = {'User-Agent': 'Mozilla/5.0'}
DELAY = 3

with open(INPUT_FILE, "r", encoding="utf-8") as f:
    all_urls = [line.strip() for line in f if line.strip()]

docx_urls = [u for u in all_urls if u.lower().endswith(".docx")]

print(f"Found {len(docx_urls)} .docx URLs to process.\n")

success_count = 0
failed_urls = []

with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
    for i, url in enumerate(docx_urls, 1):
        print(f"[{i}/{len(docx_urls)}] Fetching: {url}")
        try:
            req = urllib.request.Request(url, headers=HEADERS)
            response = urllib.request.urlopen(req, timeout=15)
            file_bytes = response.read()
        except Exception as e:
            print(f"  [FAIL download] {e}")
            failed_urls.append(url)
            time.sleep(DELAY)
            continue

        try:
            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)

            tables_text = []
            for table in document.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables_text.append("\n".join(rows))
        except Exception as e:
            print(f"  [FAIL parse] {e}")
            failed_urls.append(url)
            time.sleep(DELAY)
            continue

        record = {
            "url": url,
            "title": url.split("/")[-1],
            "text": full_text,
            "tables": tables_text,
            "scraped_at": datetime.now().isoformat()
        }

        out.write(json.dumps(record, ensure_ascii=False) + "\n")
        success_count += 1
        time.sleep(DELAY)

with open(FAILED_FILE, "w", encoding="utf-8") as f:
    for url in failed_urls:
        f.write(url + "\n")

print(f"\nDone!")
print(f"Successfully extracted: {success_count}")
print(f"Failed: {len(failed_urls)} -> {FAILED_FILE}")
print(f"Saved to: {OUTPUT_FILE}")
